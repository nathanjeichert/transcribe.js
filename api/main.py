import os
import io
import tempfile
import time
import json
import traceback
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Depends
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, ValidationError, Field
from typing import List, Optional

from google import genai
from google.genai import types
from google.api_core import exceptions as google_exceptions
from docx import Document
from docx.shared import Inches, Pt

# --- Configuration ---
# API_KEY will be read from environment variable in production (e.g., Vercel)
API_KEY = os.getenv("GEMINI_API_KEY")
MODEL_NAME = "gemini-2.5-pro-exp-03-25" # Or consider making this configurable

# --- Initialize FastAPI App ---
app = FastAPI(title="Gemini Transcriber API")

# --- CORS Middleware ---
# Allow requests from your frontend development server and production domain
# Adjust origins as needed
origins = [
    "http://localhost:3000", # Default Next.js dev port
    "http://127.0.0.1:3000",
    # Add your Vercel deployment URL(s) here later
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Initialize Gemini Client ---
# Wrap in a function to handle potential errors during startup
def get_gemini_client():
    if not API_KEY:
        print("ERROR: GEMINI_API_KEY environment variable not set.")
        # Don't raise HTTPException here, let endpoints handle it if client is None
        return None
    try:
        client = genai.Client(api_key=API_KEY)
        # Optional: Test connection (can be slow)
        # client.models.list()
        print("Gemini client initialized successfully.")
        return client
    except google_exceptions.PermissionDenied:
        print("ERROR: Permission denied initializing Gemini client. Check API key.")
        return None
    except Exception as e:
        print(f"ERROR: Failed to initialize Gemini client: {e}")
        print(traceback.format_exc())
        return None

client = get_gemini_client() # Initialize client globally

# --- Pydantic Schemas ---
class TranscriptTurn(BaseModel):
    speaker: str
    text: str

class TranscriptionRequestData(BaseModel):
    case_name: Optional[str] = None
    case_number: Optional[str] = None
    firm_name: Optional[str] = None
    input_date: Optional[str] = None # Keep as string for simplicity
    input_time: Optional[str] = None # Keep as string for simplicity
    location: Optional[str] = None
    speaker_names: Optional[List[str]] = None # List of speaker identifiers

class TranscriptionResponse(BaseModel):
    transcript_turns: List[TranscriptTurn]
    gemini_file_name: str # To identify the file for DOCX generation

class DocxRequest(BaseModel):
    gemini_file_name: str # Use the name returned by /transcribe
    title_data: dict # Contains case info, file name, duration etc.
    transcript_turns: List[TranscriptTurn]


# --- Helper Functions ---

def get_audio_mime_type(file_extension):
    """Maps file extension to MIME type for Gemini File API."""
    mime_map = {
        "mp3": "audio/mp3",
        "wav": "audio/wav",
        "aiff": "audio/aiff",
        "aac": "audio/aac",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
        # Add more mappings if needed
    }
    return mime_map.get(file_extension.lower(), None)

async def upload_to_gemini(file_path: str, mime_type: str, gemini_client: genai.Client):
    """Uploads a file to the Gemini File API and handles potential errors."""
    if not gemini_client:
        raise HTTPException(status_code=503, detail="Gemini client not available.")

    print(f"Uploading {os.path.basename(file_path)} to Gemini...")
    try:
        gemini_file = gemini_client.files.upload(file=file_path)
        print(f"File uploaded: {gemini_file.name}, processing...")

        # Poll for file readiness
        file_state = "PROCESSING"
        retries = 15
        sleep_time = 8
        max_sleep = 45
        print(f"Polling file status (up to {retries} times)...")
        while file_state == "PROCESSING" and retries > 0:
            time.sleep(sleep_time)
            try:
                file_info = gemini_client.files.get(name=gemini_file.name)
            except Exception as get_err:
                 print(f"Warning: Error checking file status (will retry): {get_err}")
                 retries -=1
                 sleep_time = min(sleep_time * 1.5, max_sleep)
                 continue

            file_state = file_info.state.name
            print(f"File state: {file_state} (retries left: {retries})")
            retries -= 1
            sleep_time = min(sleep_time * 1.5, max_sleep)

        if file_state != "ACTIVE":
            print(f"Error: File processing failed or timed out. Final state: {file_state}")
            # Attempt to delete the failed file
            try:
                gemini_client.files.delete(name=gemini_file.name)
                print(f"Warning: Attempted to delete failed/timed-out file: {gemini_file.name}")
            except Exception as del_e:
                print(f"Warning: Could not delete failed/timed-out file {gemini_file.name}: {del_e}")
            raise HTTPException(status_code=500, detail=f"Gemini file processing failed or timed out (State: {file_state})")

        print("File uploaded and processed successfully.")
        return gemini_file
    except google_exceptions.PermissionDenied:
        print("Error: Permission denied during file upload.")
        raise HTTPException(status_code=403, detail="Permission denied during Gemini file upload. Check API key permissions.")
    except google_exceptions.ResourceExhausted:
         print("Error: Resource exhausted during file upload.")
         raise HTTPException(status_code=429, detail="Gemini resource exhausted during file upload. Try again later or delete unused files.")
    except HTTPException as http_exc: # Re-raise existing HTTP exceptions
        raise http_exc
    except Exception as e:
        print(f"Error: An unexpected error occurred during file upload: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during file upload: {e}")

async def generate_transcript(gemini_file: types.File, speaker_name_list: Optional[List[str]], gemini_client: genai.Client):
    """Generates transcript using the Gemini model, requesting structured JSON output."""
    if not gemini_client:
        raise HTTPException(status_code=503, detail="Gemini client not available.")

    print("Generating transcript (requesting JSON)...")

    safety_settings=[
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HARASSMENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT, threshold=types.HarmBlockThreshold.BLOCK_NONE),
        types.SafetySetting(category=types.HarmCategory.HARM_CATEGORY_CIVIC_INTEGRITY, threshold=types.HarmBlockThreshold.BLOCK_NONE),
    ]

    if speaker_name_list:
        speaker_prompt_part = f"The speakers are identified as: {', '.join(speaker_name_list)}."
        num_speakers_part = f"There are {len(speaker_name_list)} speakers."
    else:
        speaker_prompt_part = "Speaker identifiers are not provided; use generic identifiers like SPEAKER 1, SPEAKER 2, etc., IN ALL CAPS."
        num_speakers_part = "Determine the number of speakers from the audio."

    prompt = (
        f"Generate a transcript of the speech. {num_speakers_part} {speaker_prompt_part} "
        "Structure the output STRICTLY as a JSON list of objects. "
        "Each object represents a continuous block of speech from a single speaker and MUST contain BOTH a 'speaker' field "
        "(using the provided identifiers IN ALL CAPS if available, otherwise generic ones like SPEAKER 1, SPEAKER 2, etc., IN ALL CAPS) "
        "and a 'text' field containing ALL consecutive speech from that speaker before the speaker changes. "
        "DO NOT create a new JSON object unless the speaker changes. Ensure every object has both 'speaker' and 'text' fields."
    )

    contents=[prompt, gemini_file]
    print(f"Prompt being sent to Gemini (excluding file): {prompt}")

    try:
        response = gemini_client.models.generate_content(
            model=MODEL_NAME,
            contents=contents,
            config=types.GenerateContentConfig(
                safety_settings=safety_settings,
                response_mime_type="application/json",
                response_schema=list[TranscriptTurn]
            ),
        )

        # print(f"Raw Gemini Response Text: {response.text}") # Debugging

        try:
            transcript_data = json.loads(response.text)
            validated_turns = []
            for i, turn_data in enumerate(transcript_data):
                 if 'speaker' not in turn_data:
                      print(f"Warning: Skipping turn {i+1}: Missing 'speaker' field. Data: {turn_data}")
                      continue
                 if 'text' not in turn_data:
                      print(f"Warning: Turn {i+1} for speaker '{turn_data['speaker']}' missing 'text'. Adding empty string.")
                      turn_data['text'] = ""
                 try:
                      validated_turns.append(TranscriptTurn(**turn_data))
                 except ValidationError as val_err:
                      print(f"Error: Validation error on turn {i+1}: {val_err}. Data: {turn_data}")
                      # Optionally skip or raise specific error
                      continue # Skip invalid turn

            print("Transcript generated and parsed successfully (JSON).")
            return validated_turns
        except (json.JSONDecodeError, TypeError) as json_error:
                 print(f"Error: Failed to parse JSON structure from Gemini: {json_error}")
                 print(f"Received text: {response.text}")
                 raise HTTPException(status_code=500, detail="Failed to parse transcript JSON from Gemini.")
        except AttributeError:
                 print("Error: Could not find JSON data in the expected response structure.")
                 print(f"Full response object: {response}")
                 raise HTTPException(status_code=500, detail="Could not find transcript JSON in Gemini response.")

    except google_exceptions.PermissionDenied:
        print("Error: Permission denied during transcript generation.")
        raise HTTPException(status_code=403, detail="Permission denied during Gemini transcript generation.")
    except google_exceptions.ResourceExhausted:
         print("Error: Resource exhausted during transcript generation.")
         raise HTTPException(status_code=429, detail="Gemini resource exhausted during transcript generation.")
    except HTTPException as http_exc: # Re-raise existing HTTP exceptions
        raise http_exc
    except Exception as e:
        print(f"Error: An error occurred during transcript generation: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during transcript generation: {e}")

def replace_placeholder_text(element, placeholder, replacement):
    """Replaces placeholder text in paragraphs and runs within an element (doc, cell, etc.)."""
    if hasattr(element, 'paragraphs'):
        for p in element.paragraphs:
            replace_placeholder_text(p, placeholder, replacement)
    if hasattr(element, 'runs'):
        if placeholder in element.text:
            inline = element.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    text = inline[i].text.replace(placeholder, replacement)
                    inline[i].text = text
    if hasattr(element, 'tables'):
         for table in element.tables:
             for row in table.rows:
                 for cell in row.cells:
                     replace_placeholder_text(cell, placeholder, replacement)

def create_docx(title_data: dict, transcript_turns: List[TranscriptTurn]):
    """Creates a Word document (.docx) from the transcript turns using a template."""
    template_path = "api/transcript_template.docx" # Relative to project root where FastAPI runs
    try:
        if not os.path.exists(template_path):
             print(f"Error: Template file not found at {template_path}")
             raise HTTPException(status_code=500, detail=f"Template file '{os.path.basename(template_path)}' not found on server.")
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template '{template_path}': {e}")
        raise HTTPException(status_code=500, detail=f"Error loading DOCX template: {e}")

    # Replace Title Placeholders
    for key, value in title_data.items():
        placeholder = f"{{{{{key}}}}}"
        replace_placeholder_text(doc, placeholder, str(value) if value else "")

    # Replace Transcript Body
    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if body_placeholder in p.text:
            placeholder_paragraph = p
            break

    if placeholder_paragraph:
        p_element = placeholder_paragraph._element
        p_element.getparent().remove(p_element)

        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)

            speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
            speaker_run.font.name = "Courier New"

            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        print(f"Warning: Placeholder '{body_placeholder}' not found in the template. Transcript will not be inserted.")
        # Optionally, append to end or raise an error if placeholder is mandatory
        # raise HTTPException(status_code=500, detail=f"Placeholder '{body_placeholder}' not found in template.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

async def cleanup_gemini_file(name: str, gemini_client: genai.Client):
    """Attempts to delete a file from Gemini File API."""
    if not gemini_client:
        print("Warning: Cannot cleanup Gemini file, client not available.")
        return
    if not name:
        return
    try:
        print(f"Cleaning up uploaded file: {name}...")
        gemini_client.files.delete(name=name)
        print("Cleanup successful.")
    except Exception as e:
        print(f"Warning: Could not delete uploaded file {name}: {e}")


# --- API Endpoints ---

@app.post("/transcribe", response_model=TranscriptionResponse)
async def transcribe_audio(
    request_data_json: str = Form(...), # Receive JSON data as a string field
    audio_file: UploadFile = File(...),
    gemini_client: genai.Client = Depends(get_gemini_client) # Dependency injection
):
    """
    Receives audio file and metadata, uploads to Gemini,
    generates transcript, and returns transcript data.
    """
    if not gemini_client:
        raise HTTPException(status_code=503, detail="Gemini service unavailable. Check API key/config.")

    # Parse the JSON data string
    try:
        request_data = TranscriptionRequestData.parse_raw(request_data_json)
    except ValidationError as e:
        raise HTTPException(status_code=422, detail=f"Invalid request data format: {e}")

    file_extension = audio_file.filename.split('.')[-1].lower()
    mime_type = get_audio_mime_type(file_extension)
    if not mime_type:
        raise HTTPException(status_code=400, detail=f"Unsupported audio file type: {file_extension}")

    gemini_file_obj = None
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_audio:
            content = await audio_file.read()
            temp_audio.write(content)
            temp_audio_path = temp_audio.name
            print(f"Temporary audio file saved at: {temp_audio_path}")

        # Upload to Gemini
        gemini_file_obj = await upload_to_gemini(temp_audio_path, mime_type, gemini_client)

        # Generate Transcript
        transcript_turns = await generate_transcript(gemini_file_obj, request_data.speaker_names, gemini_client)

        return TranscriptionResponse(
            transcript_turns=transcript_turns,
            gemini_file_name=gemini_file_obj.name # Return the Gemini file name
        )

    except HTTPException as e:
        # If upload failed but we got a gemini_file object, try to clean it up
        if gemini_file_obj:
            await cleanup_gemini_file(gemini_file_obj.name, gemini_client)
        raise e # Re-raise the exception
    except Exception as e:
        print(f"Error during transcription process: {e}")
        print(traceback.format_exc())
        # If upload succeeded but transcription failed, try to clean up
        if gemini_file_obj:
            await cleanup_gemini_file(gemini_file_obj.name, gemini_client)
        raise HTTPException(status_code=500, detail=f"An internal error occurred: {e}")
    finally:
        # Clean up temporary file
        if 'temp_audio_path' in locals() and os.path.exists(temp_audio_path):
            os.unlink(temp_audio_path)
            print(f"Temporary audio file deleted: {temp_audio_path}")
        # Note: We don't delete the Gemini file here automatically.
        # The frontend should call /cleanup endpoint if needed, or we rely on Gemini's TTL.


@app.post("/generate_docx")
async def generate_docx_endpoint(request: DocxRequest):
    """
    Generates a DOCX file based on provided transcript data and title info.
    """
    try:
        docx_buffer = create_docx(request.title_data, request.transcript_turns)
        file_name_base = request.title_data.get("FILE_NAME", "transcript")
        if '.' in file_name_base:
            file_name_base = os.path.splitext(file_name_base)[0]
        download_filename = f"{file_name_base}_transcript.docx"

        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{download_filename}\""}
        )
    except HTTPException as e:
        raise e # Re-raise validation or template errors
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX file: {e}")


@app.post("/cleanup/{gemini_file_name}")
async def cleanup_file(
    gemini_file_name: str,
    gemini_client: genai.Client = Depends(get_gemini_client) # Dependency injection
):
    """
    Explicitly deletes a file from the Gemini File API.
    """
    if not gemini_client:
        raise HTTPException(status_code=503, detail="Gemini service unavailable.")
    if not gemini_file_name or not gemini_file_name.startswith("files/"):
         raise HTTPException(status_code=400, detail="Invalid Gemini file name format.")

    await cleanup_gemini_file(gemini_file_name, gemini_client)
    return {"message": f"Attempted cleanup for file: {gemini_file_name}"}


# --- Root Endpoint (Optional) ---
@app.get("/")
async def root():
    return {"message": "Gemini Transcriber API is running."}

# --- Run with Uvicorn (for local development) ---
# In production (Vercel), Vercel runs the app, not this block.
if __name__ == "__main__":
    import uvicorn
    print("Starting Uvicorn server for local development...")
    # Check for API key locally
    if not API_KEY:
        print("\n*** WARNING: GEMINI_API_KEY environment variable not found. API calls will fail. ***\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
