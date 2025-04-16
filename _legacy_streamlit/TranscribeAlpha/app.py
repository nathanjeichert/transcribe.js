import streamlit as st
from google import genai
from google.genai import types
from google.api_core import exceptions as google_exceptions
from docx import Document
from docx.shared import Inches, Pt
import ffmpeg
import os
import io
import tempfile
import time
from pydub import AudioSegment
from pydantic import BaseModel, ValidationError
import json
import traceback # For detailed error logging
# --- Configuration ---
# API_KEY is now read from Streamlit secrets
MODEL_NAME = "gemini-2.5-pro-exp-03-25"
SUPPORTED_VIDEO_TYPES = ["mp4", "mov", "avi", "mkv"]
SUPPORTED_AUDIO_TYPES = ["mp3", "wav", "m4a", "flac", "ogg", "aac", "aiff"] # Add more if needed by Gemini/pydub


# --- Global Configuration ---
# Load API key from Streamlit secrets
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("GEMINI_API_KEY not found in Streamlit secrets. Please add it to your secrets file or environment variables.")
    st.stop()

# --- Initialize Gemini Client ---
try:
    client = genai.Client(api_key=API_KEY)
    # Optional: Test configuration by listing models (can be slow)
    # client.models.list()
except google_exceptions.PermissionDenied:
    st.error("Permission denied initializing Gemini client. Please check your API key (loaded from secrets).")
    st.stop() # Stop execution if initialization fails
except Exception as e:
    st.error(f"Failed to initialize Gemini client: {e}")
    st.error(traceback.format_exc()) # Log traceback for debugging
    st.stop() # Stop execution if initialization fails


# --- Pydantic Schema ---
class TranscriptTurn(BaseModel):
  speaker: str
  text: str # Renamed from dialogue
  
def convert_video_to_audio(input_path, output_path, format="mp3"):
    """Converts video file to audio using ffmpeg."""
    try:
        st.info(f"Converting video to {format} audio...")
        (
            ffmpeg
            .input(input_path)
            .output(output_path, format=format, acodec='libmp3lame') # Specify mp3 codec
            .overwrite_output()
            .run(quiet=True) # Suppress ffmpeg output in console
        )
        st.success("Video converted to audio successfully.")
        return output_path
    except ffmpeg.Error as e:
        st.error(f"FFmpeg error during conversion: {e.stderr.decode() if e.stderr else 'Unknown error'}")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred during video conversion: {e}")
        return None

def get_audio_mime_type(file_extension):
    """Maps file extension to MIME type for Gemini File API."""
    mime_map = {
        "mp3": "audio/mp3",
        "wav": "audio/wav",
        "aiff": "audio/aiff",
        "aac": "audio/aac",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
        # Add more mappings if needed
    }
    return mime_map.get(file_extension.lower(), None)

def upload_to_gemini(file_path, mime_type):
    """Uploads a file to the Gemini File API and handles potential errors."""
    st.info(f"Uploading {os.path.basename(file_path)} to Gemini...")
    try:
        # Use client.files.upload (mime_type is likely inferred or not needed as direct arg)
        gemini_file = client.files.upload(file=file_path) # Removed mime_type argument
        st.info(f"File uploaded: {gemini_file.name}, processing...")

        # Poll for file readiness
        file_state = "PROCESSING"
        retries = 15 # Slightly more retries
        sleep_time = 8 # Slightly longer initial sleep
        max_sleep = 45 # Max sleep time
        st.info(f"Polling file status (up to {retries} times)...")
        while file_state == "PROCESSING" and retries > 0:
            time.sleep(sleep_time) # Wait before checking status again
            try:
                # Use client.files.get
                file_info = client.files.get(name=gemini_file.name)
            except Exception as get_err:
                 st.warning(f"Error checking file status (will retry): {get_err}")
                 retries -=1
                 sleep_time = min(sleep_time * 1.5, max_sleep)
                 continue # Skip to next retry

            file_state = file_info.state.name
            st.info(f"File state: {file_state} (retries left: {retries})")
            retries -= 1
            sleep_time = min(sleep_time * 1.5, max_sleep) # Exponential backoff

        if file_state != "ACTIVE":
            st.error(f"File processing failed or timed out after polling. Final state: {file_state}")
            # Attempt to delete the failed file
            try:
                # Use client.files.delete
                client.files.delete(name=gemini_file.name)
                st.warning(f"Attempted to delete failed/timed-out file: {gemini_file.name}")
            except Exception as del_e:
                st.warning(f"Could not delete failed/timed-out file {gemini_file.name}: {del_e}")
            return None

        st.success("File uploaded and processed successfully.")
        return gemini_file
    except google_exceptions.PermissionDenied:
        st.error("Permission denied during file upload. Check API key permissions.")
        return None
    except google_exceptions.ResourceExhausted:
         st.error("Resource exhausted during file upload. You might have uploaded too many files or exceeded storage limits. Please try again later or delete unused files.")
         return None
    except Exception as e:
        st.error(f"An error occurred during file upload: {e}")
        st.error(traceback.format_exc()) # Log traceback
        return None

def generate_transcript(gemini_file, speaker_name_list=None):
    """Generates transcript using the Gemini model, always requesting structured JSON output."""

    st.info("Generating transcript (requesting JSON)...")

    # Define safety settings (remains the same structure)
    safety_settings=[
        types.SafetySetting(
            category=types.HarmCategory.HARM_CATEGORY_HARASSMENT,
            threshold=types.HarmBlockThreshold.BLOCK_NONE,
        ),
        types.SafetySetting(
            category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH,
            threshold=types.HarmBlockThreshold.BLOCK_NONE,
        ),
        types.SafetySetting(
            category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
            threshold=types.HarmBlockThreshold.BLOCK_NONE,
        ),
        types.SafetySetting(
            category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
            threshold=types.HarmBlockThreshold.BLOCK_NONE,
        ),
         types.SafetySetting(
            category=types.HarmCategory.HARM_CATEGORY_CIVIC_INTEGRITY,
            threshold=types.HarmBlockThreshold.BLOCK_NONE,
        ),
    ]


    # Construct the prompt, mentioning speakers are optional
    if speaker_name_list:
        speaker_prompt_part = f"The speakers are identified as: {', '.join(speaker_name_list)}."
        num_speakers_part = f"There are {len(speaker_name_list)} speakers."
    else:
        speaker_prompt_part = "Speaker identifiers are not provided; use generic identifiers like SPEAKER 1, SPEAKER 2, etc., IN ALL CAPS." # Removed underscore
        num_speakers_part = "Determine the number of speakers from the audio."

    # Ensure this block is indented to match the 'if speaker_name_list:' line above
    prompt = (
        f"Generate a transcript of the speech. {num_speakers_part} {speaker_prompt_part} "
        "Structure the output STRICTLY as a JSON list of objects. "
            "Each object represents a continuous block of speech from a single speaker and MUST contain BOTH a 'speaker' field "
            "(using the provided identifiers IN ALL CAPS if available, otherwise generic ones like SPEAKER_1, SPEAKER_2, etc., IN ALL CAPS) "
            "and a 'text' field containing ALL consecutive speech from that speaker before the speaker changes. "
            "DO NOT create a new JSON object unless the speaker changes. Ensure every object has both 'speaker' and 'text' fields."

        )

    # Define contents for the request
    contents=[prompt, gemini_file]

    st.caption(f"Prompt being sent to Gemini (excluding file): {prompt}") # Show the updated prompt

    try:
        # Use client.models.generate_content
        # Pass all config settings directly within a single GenerateContentConfig object
        response = client.models.generate_content(
            model=MODEL_NAME, # Specify model here
            contents=contents,
            config=types.GenerateContentConfig(
                safety_settings=safety_settings,
                response_mime_type="application/json", # Moved from nested config
                response_schema=list[TranscriptTurn]  # Moved from nested config
            ),
        )

        # Debug: Print raw response text if needed
        # st.write("Raw Gemini Response Text:")
        # st.code(f"Raw response text: {response.text}")

        # Check for empty or blocked response FIRST
        # Handle potential lack of 'parts' if blocked early
        # Always process as JSON now
        try:
            transcript_data = json.loads(response.text)
            # Initial validation might still fail if 'text' is missing, handle it
            validated_turns = []
            for i, turn_data in enumerate(transcript_data):
                 # Ensure both fields exist, provide default for 'text' if missing initially
                 if 'speaker' not in turn_data:
                      st.warning(f"Skipping turn {i+1}: Missing 'speaker' field. Data: {turn_data}")
                      continue
                 if 'text' not in turn_data: # Check for 'text' field now
                      st.warning(f"Turn {i+1} for speaker '{turn_data['speaker']}' missing 'text'. Adding empty string.")
                      turn_data['text'] = "" # Add missing field
                 try:
                      validated_turns.append(TranscriptTurn(**turn_data))
                 except ValidationError as val_err:
                      st.error(f"Validation error on turn {i+1}: {val_err}. Data: {turn_data}")
                      # Optionally skip this turn or handle differently
                      continue

            # No merging needed now, return validated turns directly
            st.success("Transcript generated and parsed successfully (JSON).")
            return validated_turns # Return the validated list
        except (json.JSONDecodeError, TypeError) as json_error: # Handles inner try failure
                 st.error(f"Failed to parse the JSON structure from Gemini: {json_error}") # Indent this block
                 st.error("Received text (may not be valid JSON):")
                 st.code(response.text)
                 return None
        except AttributeError: # Align this except with the previous one
                 st.error("Could not find JSON data in the expected response structure.") # Indent this block
                 st.write(f"Full response object: {response}")
                 return None
        # This except block corresponds to the outer try block starting before model.generate_content
    except google_exceptions.PermissionDenied:
        st.error("Permission denied during transcript generation. Check API key permissions.")
        st.error(traceback.format_exc())
        return None
    except google_exceptions.ResourceExhausted:
         st.error("Resource exhausted during transcript generation. The model might be overloaded or the input too large. Please try again later.")
         st.error(traceback.format_exc())
         return None
    except Exception as e:
        st.error(f"An error occurred during transcript generation: {e}")
        st.error(traceback.format_exc()) # Log traceback
        # Consider adding more specific error handling based on potential Gemini API errors
        return None

def replace_placeholder_text(element, placeholder, replacement):
    """Replaces placeholder text in paragraphs and runs within an element (doc, cell, etc.)."""
    if hasattr(element, 'paragraphs'):
        for p in element.paragraphs:
            replace_placeholder_text(p, placeholder, replacement) # Recurse for paragraphs
    if hasattr(element, 'runs'):
        if placeholder in element.text:
            inline = element.runs
            # Replace strings and retain formatting
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    text = inline[i].text.replace(placeholder, replacement)
                    inline[i].text = text
    if hasattr(element, 'tables'):
         for table in element.tables:
             for row in table.rows:
                 for cell in row.cells:
                     replace_placeholder_text(cell, placeholder, replacement) # Recurse for cells

def create_docx(title_data, transcript_turns):
    """Creates a Word document (.docx) from the transcript turns (list of TranscriptTurn) using a template."""
    try:
        doc = Document("transcript_template.docx") # Load the template
    except Exception as e:
        st.error(f"Error loading template 'gemini_transcriber/transcript_template.docx': {e}")
        return None

    # --- Replace Title Placeholders ---
    for key, value in title_data.items():
        placeholder = f"{{{{{key}}}}}" # e.g., {{CASE_NAME}}
        replace_placeholder_text(doc, placeholder, str(value) if value else "") # Replace in whole doc

    # --- Replace Transcript Body ---
    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if body_placeholder in p.text:
            placeholder_paragraph = p
            break # Found the placeholder

    if placeholder_paragraph:
        # Remove the placeholder paragraph (we'll insert new ones)
        p_element = placeholder_paragraph._element
        p_element.getparent().remove(p_element)
        
        # Add each turn as a new paragraph
        for turn in transcript_turns:
            # Create a new paragraph
            p = doc.add_paragraph()
            
            # Format paragraph
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            
            # Add speaker and text with Courier New font
            speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
            speaker_run.font.name = "Courier New"
            
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        st.warning(f"Placeholder '{body_placeholder}' not found in the template. Appending transcript to the end.")
        # Fallback: Append to end if placeholder not found
        for turn in transcript_turns:
            p = doc.add_paragraph()
            
            # Format paragraph
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            
            # Add speaker and text with Courier New font
            speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
            speaker_run.font.name = "Courier New"
            
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit App ---

st.set_page_config(layout="wide")
st.title("📄 Gemini Legal Transcript Generator")
st.markdown("Upload an audio or video file to generate a transcript using Gemini 2.5 Pro.")

# --- Input Fields ---
st.subheader("Case Information (Optional)")
col1, col2 = st.columns(2)
with col1:
    case_name = st.text_input("Case Name:")
    case_number = st.text_input("Case Number:")
    firm_name = st.text_input("Firm or Organization Name:")
with col2:
    input_date = st.date_input("Date:")
    input_time = st.time_input("Time:")
    location = st.text_input("Location:")

st.subheader("Speaker Information (Optional)")
specify_speakers = st.checkbox("Manually specify speaker names?")
speaker_names_input = None # Default to None (auto-detect/generic)
if specify_speakers:
    num_speakers = st.number_input("Number of Speakers:", min_value=1, value=2, step=1, key="num_speakers_input")
    speaker_names_input = [] # Initialize list only if checkbox is ticked
    for i in range(num_speakers):
        speaker_name = st.text_input(f"Speaker {i+1} Identifier (ALL CAPS):", key=f"speaker_{i}", help="Enter the exact identifier the model should use (e.g., COUNSEL, WITNESS, JOHN_DOE).")
        # Use entered name directly, ensure it's uppercase, or default if empty (no underscore)
        speaker_names_input.append(speaker_name.strip().upper() if speaker_name else f"SPEAKER {i+1}")

st.markdown("---") # Separator

uploaded_file = st.file_uploader(
    "Choose an audio or video file for Transcription", # Changed label slightly
        type=SUPPORTED_AUDIO_TYPES + SUPPORTED_VIDEO_TYPES,
        accept_multiple_files=False
    )

# Store uploaded file info in session state to persist across button clicks
if uploaded_file is not None:
    st.session_state.uploaded_file_info = {
        "name": uploaded_file.name,
        "type": uploaded_file.type,
        "size": uploaded_file.size,
        "bytes": uploaded_file.getvalue() # Read bytes here
    }
    st.write(f"Uploaded file: `{st.session_state.uploaded_file_info['name']}` ({st.session_state.uploaded_file_info['size']} bytes)")
    # Display the "Generate" button only after a file is uploaded
    generate_button = st.button("Generate Transcript", type="primary")
else:
    # Clear session state if no file is uploaded
    if 'uploaded_file_info' in st.session_state:
        del st.session_state.uploaded_file_info
    generate_button = False # Ensure button doesn't trigger anything

# --- Processing Triggered by Button ---
if generate_button and 'uploaded_file_info' in st.session_state:
    with st.spinner("Processing file and generating transcript... Please wait."):
        # Retrieve file info from session state
        file_info = st.session_state.uploaded_file_info
        file_name = file_info['name']
        file_extension = file_name.split('.')[-1].lower()
        file_bytes = file_info['bytes']

        # Prepare title data dictionary
        title_data = {
            "CASE_NAME": case_name,
            "CASE_NUMBER": case_number,
            "FIRM_OR_ORGANIZATION_NAME": firm_name,
            "DATE": str(input_date) if input_date else "",
            "TIME": str(input_time) if input_time else "",
            "LOCATION": location,
            "FILE_NAME": file_name, # Auto-populated
            "FILE_DURATION": "Calculating..." # Placeholder
        }

        # Use a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, file_name)
            # Write bytes from session state to temp file
            with open(input_path, "wb") as f:
                f.write(file_bytes)

            audio_path = None
            mime_type = None
            gemini_file_to_delete = None # Keep track of the file to delete
            file_duration_str = "N/A" # Default duration

            try:
                # --- File Processing & Duration Calculation ---
                if file_extension in SUPPORTED_VIDEO_TYPES:
                    st.write("Video file detected. Converting to audio...")
                    output_audio_filename = f"{os.path.splitext(file_name)[0]}.mp3"
                    output_path = os.path.join(temp_dir, output_audio_filename)
                    audio_path = convert_video_to_audio(input_path, output_path, format="mp3")
                    if audio_path:
                        mime_type = get_audio_mime_type("mp3")
                elif file_extension in SUPPORTED_AUDIO_TYPES:
                    st.write("Audio file detected.") # Correct indentation
                    audio_path = input_path
                    mime_type = get_audio_mime_type(file_extension)
                    # Handle potential conversion if needed (e.g., m4a to mp3 if Gemini prefers)
                    if file_extension == 'm4a': # Example: Convert m4a as pydub handles it well
                        st.info("Converting M4A to MP3 for better compatibility...")
                        try:
                            sound = AudioSegment.from_file(audio_path, format="m4a")
                            output_audio_filename = f"{os.path.splitext(file_name)[0]}.mp3"
                            output_path = os.path.join(temp_dir, output_audio_filename)
                            sound.export(output_path, format="mp3")
                            audio_path = output_path
                            mime_type = get_audio_mime_type("mp3")
                            st.success("M4A converted to MP3.")
                        except Exception as convert_err:
                            st.error(f"Could not convert M4A: {convert_err}")
                            audio_path = None # Prevent upload if conversion failed
                else:
                    st.error(f"Unsupported file type: {file_extension}")

                # Calculate duration if we have a valid audio path
                if audio_path:
                    try:
                        audio_segment = AudioSegment.from_file(audio_path)
                        duration_seconds = len(audio_segment) / 1000
                        # Format duration HH:MM:SS
                        hours, rem = divmod(duration_seconds, 3600)
                        minutes, seconds = divmod(rem, 60)
                        file_duration_str = "{:0>2}:{:0>2}:{:0>2}".format(int(hours), int(minutes), int(round(seconds)))
                        title_data["FILE_DURATION"] = file_duration_str # Update dict
                        st.write(f"Audio Duration: {file_duration_str}")
                    except Exception as e:
                        st.warning(f"Could not determine audio duration: {e}")
                        title_data["FILE_DURATION"] = "N/A"


                # --- Gemini Interaction ---
                if audio_path and mime_type:
                    # Pass file path and mime type directly
                    gemini_file = upload_to_gemini(audio_path, mime_type)
                    gemini_file_to_delete = gemini_file # Mark for deletion

                    if gemini_file:
                        # Pass the uploaded gemini_file object and speaker names (which might be None)
                        transcript_content = generate_transcript(gemini_file, speaker_names_input)

                        # transcript_content should now always be a list (or None if failed)
                        if isinstance(transcript_content, list):
                            st.subheader("Generated Transcript:")

                            # Display structured transcript in text area
                            full_transcript_text = "\n\n".join([f"{turn.speaker.upper()}:\t\t{turn.text}" for turn in transcript_content]) # Use text field
                            st.text_area("Transcript Preview", full_transcript_text, height=400)

                            # --- Download Button ---
                            docx_buffer = create_docx(title_data, transcript_content) # Pass title data and the list of turns
                            st.download_button(
                                label="Download Transcript (.docx)",
                                data=docx_buffer,
                                file_name=f"{os.path.splitext(file_name)[0]}_transcript.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.warning("Could not generate transcript.")
                    else:
                        st.warning("Could not upload file to Gemini.")
                elif audio_path and not mime_type: # Ensure block below is indented
                    st.error(f"Could not determine MIME type for audio format: {file_extension}")

            finally:
                # --- Cleanup ---
                if gemini_file_to_delete:
                    try:
                        st.info(f"Cleaning up uploaded file: {gemini_file_to_delete.name}...")
                        # Use client.files.delete directly
                        client.files.delete(name=gemini_file_to_delete.name)
                        st.success("Cleanup successful.")
                    except Exception as e:
                        st.warning(f"Could not delete uploaded file {gemini_file_to_delete.name}: {e}")
                # Temporary directory is automatically cleaned up by 'with' statement

# Ensure these lines are definitely at the top level (no indentation)
st.markdown("---")
st.caption("Powered by Google Gemini")
